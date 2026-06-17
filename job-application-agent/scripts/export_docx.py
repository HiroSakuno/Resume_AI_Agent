#!/usr/bin/env python3
"""
Export Markdown files to .docx using python-docx.
Called by generate_application.py.

Usage:
    python scripts/export_docx.py <resume.md> <cover_letter.md> <output_dir>
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_markdown_to_docx(md_text: str, doc: Document, is_cover_letter: bool = False):
    """
    Convert markdown text into a Word document.
    Handles: H1, H2, H3, bullet lists, bold, horizontal rules, plain paragraphs.
    """
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines (add spacing)
        if not line.strip():
            i += 1
            continue

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.color.rgb = RGBColor(0x1F, 0x29, 0x7A)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            p.runs[0].font.size = Pt(13)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            p.runs[0].font.size = Pt(11)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph("_" * 60)
            i += 1
            continue

        # Bullet point
        if line.startswith("- ") or line.startswith("* "):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, bullet_text)
            i += 1
            continue

        # Plain paragraph (apply inline formatting)
        p = doc.add_paragraph()
        _add_formatted_run(p, line.strip())
        i += 1

    return doc


def _add_formatted_run(paragraph, text: str):
    """
    Add text to a paragraph, handling inline **bold** and *italic* markdown.
    """
    # Split on bold (**text**) and italic (*text*)
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def set_document_styles(doc: Document):
    """Set global document styles for ATS-friendly formatting."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def export_md_to_docx(md_path: Path, docx_path: Path, is_cover_letter: bool = False):
    """Convert a markdown file to a .docx file."""
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()
    set_document_styles(doc)
    parse_markdown_to_docx(md_text, doc, is_cover_letter)
    doc.save(str(docx_path))
    print(f"  Exported: {docx_path.name}")


def main():
    if len(sys.argv) < 4:
        print("Usage: export_docx.py <resume.md> <cover_letter.md> <output_dir>")
        sys.exit(1)

    resume_md = Path(sys.argv[1])
    cover_md = Path(sys.argv[2])
    output_dir = Path(sys.argv[3])

    if resume_md.exists():
        export_md_to_docx(resume_md, output_dir / "resume.docx", is_cover_letter=False)
    else:
        print(f"  WARNING: {resume_md} not found, skipping DOCX export for resume.")

    if cover_md.exists():
        export_md_to_docx(cover_md, output_dir / "cover_letter.docx", is_cover_letter=True)
    else:
        print(f"  WARNING: {cover_md} not found, skipping DOCX export for cover letter.")


if __name__ == "__main__":
    main()
